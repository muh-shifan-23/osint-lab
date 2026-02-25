from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import PyPDF2
import docx
from datetime import datetime
import os
import traceback

class MetadataExtractor:
    def __init__(self):
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
            'document': ['.pdf', '.docx', '.doc'],
            'other': ['.mp3', '.mp4']
        }
    
    def extract(self, file_path):
        """Extract metadata from file"""
        ext = os.path.splitext(file_path)[1].lower()
        
        result = {
            'filename': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'extension': ext,
            'timestamp': datetime.now().isoformat(),
            'metadata': {}
        }
        
        try:
            if ext in self.supported_formats['image']:
                result['metadata'] = self.extract_image_metadata(file_path)
            elif ext == '.pdf':
                result['metadata'] = self.extract_pdf_metadata(file_path)
            elif ext in ['.docx', '.doc']:
                result['metadata'] = self.extract_docx_metadata(file_path)
            else:
                result['metadata'] = {'error': 'Unsupported file format'}
        except Exception as e:
            result['metadata'] = {
                'error': str(e),
                'error_type': type(e).__name__,
                'traceback': traceback.format_exc()
            }
        
        return result
    
    def extract_image_metadata(self, file_path):
        """Extract EXIF data from images"""
        exif_data = {}
        
        try:
            image = Image.open(file_path)
            
            # Basic image info
            exif_data['format'] = str(image.format) if image.format else 'Unknown'
            exif_data['mode'] = str(image.mode) if image.mode else 'Unknown'
            exif_data['size'] = {
                'width': image.size[0],
                'height': image.size[1]
            }
            
            # Try to extract EXIF data
            try:
                exif = image._getexif()
                if exif and isinstance(exif, dict):
                    exif_data['exif'] = {}
                    
                    for tag_id, value in exif.items():
                        try:
                            # Get tag name
                            tag = TAGS.get(tag_id, f'Unknown_{tag_id}')
                            
                            # Handle GPS data separately
                            if tag == 'GPSInfo' and isinstance(value, dict):
                                gps_data = self.parse_gps_info(value)
                                if gps_data:
                                    exif_data['exif']['GPSInfo'] = gps_data
                                    
                                    # Try to get coordinates
                                    coords = self.extract_coordinates(gps_data)
                                    if coords:
                                        exif_data['coordinates'] = coords
                            else:
                                # Convert value to safe string
                                exif_data['exif'][tag] = self.safe_convert_value(value)
                                
                        except Exception as tag_error:
                            # Skip problematic tags
                            continue
                else:
                    exif_data['exif'] = {'note': 'No EXIF data found in image'}
                    
            except AttributeError:
                exif_data['exif'] = {'note': 'Image format does not support EXIF data'}
            except Exception as exif_error:
                exif_data['exif'] = {'error': f'Error reading EXIF: {str(exif_error)}'}
            
            return exif_data
            
        except Exception as e:
            return {
                'error': str(e),
                'error_type': type(e).__name__
            }
    
    def safe_convert_value(self, value):
        """Safely convert any value to a JSON-serializable format"""
        try:
            # Handle bytes
            if isinstance(value, bytes):
                try:
                    return value.decode('utf-8', errors='ignore')
                except:
                    return f'<binary data, {len(value)} bytes>'
            
            # Handle tuples and lists
            elif isinstance(value, (tuple, list)):
                try:
                    # Try to convert to simple list
                    converted = []
                    for item in value:
                        if isinstance(item, (int, float, str)):
                            converted.append(item)
                        elif isinstance(item, tuple) and len(item) == 2:
                            # Handle rational numbers (numerator, denominator)
                            converted.append(f"{item[0]}/{item[1]}")
                        else:
                            converted.append(str(item))
                    return converted
                except:
                    return str(value)
            
            # Handle dictionaries
            elif isinstance(value, dict):
                return {k: self.safe_convert_value(v) for k, v in value.items()}
            
            # Handle basic types
            elif isinstance(value, (int, float, str, bool, type(None))):
                return value
            
            # Everything else - convert to string
            else:
                return str(value)
                
        except Exception as e:
            return f'<conversion error: {str(e)}>'
    
    def parse_gps_info(self, gps_info):
        """Parse GPS information from EXIF"""
        parsed = {}
        
        try:
            for gps_tag, value in gps_info.items():
                try:
                    tag_name = GPSTAGS.get(gps_tag, f'Unknown_{gps_tag}')
                    parsed[tag_name] = self.safe_convert_value(value)
                except:
                    continue
        except:
            pass
        
        return parsed
    
    def extract_coordinates(self, gps_data):
        """Extract and convert GPS coordinates to decimal degrees"""
        try:
            # Get required GPS fields
            lat = gps_data.get('GPSLatitude')
            lat_ref = gps_data.get('GPSLatitudeRef')
            lon = gps_data.get('GPSLongitude')
            lon_ref = gps_data.get('GPSLongitudeRef')
            
            if not all([lat, lat_ref, lon, lon_ref]):
                return None
            
            # Convert to decimal degrees
            lat_decimal = self.convert_to_degrees(lat, lat_ref)
            lon_decimal = self.convert_to_degrees(lon, lon_ref)
            
            if lat_decimal is not None and lon_decimal is not None:
                return {
                    'latitude': lat_decimal,
                    'longitude': lon_decimal,
                    'google_maps': f"https://www.google.com/maps?q={lat_decimal},{lon_decimal}"
                }
        except:
            pass
        
        return None
    
    def convert_to_degrees(self, value, ref):
        """Convert GPS coordinates to decimal degrees"""
        try:
            # Handle different value formats
            if isinstance(value, str):
                # Already converted - try to parse
                parts = value.replace('[', '').replace(']', '').split(',')
                if len(parts) == 3:
                    d = float(parts[0].strip())
                    m = float(parts[1].strip())
                    s = float(parts[2].strip())
                else:
                    return None
            elif isinstance(value, (list, tuple)) and len(value) == 3:
                # Convert each component
                d = self.parse_coordinate_part(value[0])
                m = self.parse_coordinate_part(value[1])
                s = self.parse_coordinate_part(value[2])
                
                if None in [d, m, s]:
                    return None
            else:
                return None
            
            # Calculate decimal degrees
            decimal = d + (m / 60.0) + (s / 3600.0)
            
            # Apply reference (N/S for latitude, E/W for longitude)
            if ref in ['S', 'W']:
                decimal = -decimal
            
            return decimal
            
        except Exception as e:
            return None
    
    def parse_coordinate_part(self, part):
        """Parse a single coordinate part (degrees, minutes, or seconds)"""
        try:
            if isinstance(part, (int, float)):
                return float(part)
            elif isinstance(part, str):
                # Handle fraction format "123/456"
                if '/' in part:
                    num, den = part.split('/')
                    return float(num) / float(den)
                return float(part)
            elif isinstance(part, tuple) and len(part) == 2:
                # Handle tuple format (numerator, denominator)
                return float(part[0]) / float(part[1])
            else:
                return None
        except:
            return None
    
    def extract_pdf_metadata(self, file_path):
        """Extract metadata from PDF files"""
        try:
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                info = pdf.metadata
                
                metadata = {
                    'pages': len(pdf.pages),
                    'encrypted': pdf.is_encrypted
                }
                
                if info:
                    metadata['author'] = str(info.get('/Author', 'N/A'))
                    metadata['creator'] = str(info.get('/Creator', 'N/A'))
                    metadata['producer'] = str(info.get('/Producer', 'N/A'))
                    metadata['subject'] = str(info.get('/Subject', 'N/A'))
                    metadata['title'] = str(info.get('/Title', 'N/A'))
                    metadata['creation_date'] = str(info.get('/CreationDate', 'N/A'))
                    metadata['modification_date'] = str(info.get('/ModDate', 'N/A'))
                
                return metadata
        except Exception as e:
            return {'error': str(e), 'error_type': type(e).__name__}
    
    def extract_docx_metadata(self, file_path):
        """Extract metadata from DOCX files"""
        try:
            doc = docx.Document(file_path)
            core_properties = doc.core_properties
            
            metadata = {
                'author': str(core_properties.author) if core_properties.author else 'N/A',
                'category': str(core_properties.category) if core_properties.category else 'N/A',
                'comments': str(core_properties.comments) if core_properties.comments else 'N/A',
                'content_status': str(core_properties.content_status) if core_properties.content_status else 'N/A',
                'created': str(core_properties.created) if core_properties.created else 'N/A',
                'identifier': str(core_properties.identifier) if core_properties.identifier else 'N/A',
                'keywords': str(core_properties.keywords) if core_properties.keywords else 'N/A',
                'language': str(core_properties.language) if core_properties.language else 'N/A',
                'last_modified_by': str(core_properties.last_modified_by) if core_properties.last_modified_by else 'N/A',
                'last_printed': str(core_properties.last_printed) if core_properties.last_printed else 'N/A',
                'modified': str(core_properties.modified) if core_properties.modified else 'N/A',
                'revision': str(core_properties.revision) if core_properties.revision else 'N/A',
                'subject': str(core_properties.subject) if core_properties.subject else 'N/A',
                'title': str(core_properties.title) if core_properties.title else 'N/A',
                'version': str(core_properties.version) if core_properties.version else 'N/A',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return metadata
        except Exception as e:
            return {'error': str(e), 'error_type': type(e).__name__}